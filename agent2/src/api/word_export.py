# -*- coding: utf-8 -*-
"""
【模块作用】将 Markdown 粗解析为 Word（标题、简单表格、加粗），供 /api/agent/export-word 使用。

【MySQL 迁移】无需修改。
"""
from __future__ import annotations  # 延迟注解

import io  # BytesIO：内存中接收 Document.save 输出
import re  # 拆分表格分隔行、粗体 **...**
from typing import List  # 表格 rows 的类型别名

from docx import Document  # python-docx：新建空白文档
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 一级标题左对齐
from docx.shared import Pt  # 字号单位：磅


def markdown_to_docx_bytes(markdown: str) -> bytes:
    """
    极简按行状态机扫描：
    - 空行跳过；
    - # / ## / ### 转为 Heading 1/2/3；
    - 以 | 开头的连续行解析为表格（跳过分隔线如 |---|）；
    - 其余行作为正文，支持行内 **粗体**。
    复杂语法（列表、链接、代码块）未实现，可换服务端 pandoc。
    """
    doc = Document()  # 新建 Word 文档对象
    style = doc.styles["Normal"]  # 修改默认正文样式
    style.font.name = "宋体"  # 中文字体
    style.font.size = Pt(11)  # 正文 11 磅

    lines = markdown.splitlines()  # 按换行拆成行列表（不含 \\n 字符）
    i = 0  # 当前扫描行下标
    while i < len(lines):
        line = lines[i].rstrip()  # 去掉行尾空白，保留行首缩进意义不大此处未专门处理
        if not line.strip():
            i += 1  # 纯空行：不占段落
            continue
        if line.startswith("# "):
            # 一级标题：去掉 "# " 前缀
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 显式左对齐
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip().startswith("|") and "|" in line:
            # 可能进入多行表格块：rows 收集所有数据行
            rows: List[List[str]] = []
            while i < len(lines) and "|" in lines[i]:
                row = [c.strip() for c in lines[i].split("|") if c.strip() != ""]
                # 跳过分隔行：仅由 - : 空格 | 组成
                if row and not re.match(r"^[-:\s|]+$", lines[i]):
                    rows.append(row)
                i += 1  # 内层循环已前进 i
            i -= 1  # 外层循环末尾还会 i+=1，此处回退一格对齐
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))  # 假定各行列数一致
                table.style = "Table Grid"  # 带边框网格样式
                for ri, r in enumerate(rows):
                    for ci, cell in enumerate(r):
                        if ci < len(table.rows[ri].cells):
                            table.rows[ri].cells[ci].text = cell
        else:
            p = doc.add_paragraph()  # 空段落，再分段添加 run
            _add_inline(p, line)
        i += 1  # 处理完当前行（或表格块）后前进

    buf = io.BytesIO()  # 内存文件句柄
    doc.save(buf)  # python-docx 写入 BytesIO
    return buf.getvalue()  # 原始 docx 字节，供 HTTP 响应


def _add_inline(paragraph, text: str) -> None:
    """
    用正则捕获组切分 **bold** 与普通文本；
    re.split 保留捕获组，故粗体片段也会出现在 parts 里。
    """
    parts = re.split(r"(\*\*[^*]+\*\*)", text)  # 奇偶交替：普通、粗体、普通…
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])  # 去掉首尾 **
            run.bold = True
        else:
            paragraph.add_run(part)  # 普通文本 run
